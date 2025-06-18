import os
import re
from docx import Document
import json
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def extract_clauses(docx_path, file_name=None):
    doc = Document(docx_path)
    clauses = []
    current_clause = {"number": "", "content": ""}
    in_clause = False
    skip_section = False

    main_clause_pattern = re.compile(r'^第(?:[一二三四五六七八九十百]+|[\d]+)条\s+', re.IGNORECASE)
    subclause_pattern = re.compile(r'^\s*(?:\(一\)|\(二\)|\(三\)|\(四\)|\(五\)|\(六\)|\(七\)|\(八\)|'
                                   r'\(1\)|\(2\)|\(3\)|\(4\)|\(5\)|\(6\)|\(7\)|\(8\)|'
                                   r'[一二三四五六七八九十]、|\d+\.\s*)\s*', re.IGNORECASE)
    section_pattern = re.compile(r'^第[一二三四五六七八九十]+章\s+', re.IGNORECASE)
    chapter_title_pattern = re.compile(r'^[一二三四五六七八九十]+、[\u4e00-\u9fa5]+$')

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if section_pattern.match(text) or chapter_title_pattern.match(text):
            if "附则" in text:
                skip_section = True
            else:
                skip_section = False
            continue

        if skip_section:
            continue

        if main_clause_pattern.match(text):
            if current_clause["content"]:
                clauses.append(current_clause)
                current_clause = {"number": "", "content": ""}

            number_match = main_clause_pattern.match(text)
            current_clause["number"] = number_match.group().strip()
            current_clause["content"] = text[number_match.end():].strip()
            in_clause = True

        elif in_clause and subclause_pattern.match(text):
            current_clause["content"] += "\n" + text

        elif in_clause:
            current_clause["content"] += "\n" + text

    if current_clause["content"] and not skip_section:
        clauses.append(current_clause)

    if file_name:
        for clause in clauses:
            clause["file_name"] = file_name

    return clauses


def save_to_word(clauses, output_path):
    doc = Document()
    title = doc.add_heading("条款提取结果", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, clause in enumerate(clauses, 1):
        file_prefix = f"{clause.get('file_name', '')}条款 " if clause.get('file_name') else "条款 "
        clause_title = doc.add_heading(f"{file_prefix}{i}: {clause['number']}", level=2)
        content_paragraph = doc.add_paragraph(clause['content'])
        doc.add_paragraph("-" * 50)

    doc.save(output_path)
    print(f"✅ 已保存Word文件至 {output_path}")


if __name__ == "__main__":
    folder_path = "rawdata"  # 文件夹名称
    all_clauses = []

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith('.docx'):
                file_path = os.path.join(root, file)
                file_name = os.path.splitext(file)[0]
                clauses = extract_clauses(file_path, file_name)
                all_clauses.extend(clauses)

    output_path = "combined_clauses.docx"
    save_to_word(all_clauses, output_path)

    print(f"📊 共提取 {len(all_clauses)} 条有效条款")