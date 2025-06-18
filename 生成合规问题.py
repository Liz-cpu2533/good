import asyncio
import os
import time
import re
from docx import Document
from tqdm import tqdm
from volcenginesdkarkruntime import AsyncArk

# 直接固定配置参数
API_KEY = "f46d245b-7dbe-48e7-ac8a-4c19740cd14a"  # 火山引擎API密钥
MODEL_ID = "ep-bi-20250418153651-h8dqv"  # V3模型ID
QUESTION_PER_CLAUSE = 3  # 每条条款生成的问题数量

# 提取条款
def extract_clauses(docx_path):
    doc = Document(docx_path)
    text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    clauses = re.split(r"[。\n；;]+", text)
    return [c.strip() for c in clauses if len(c.strip()) > 8]

# 构造prompt
def build_prompt(clause, k=3):
    return (
        f"你是一位投资者，试图绕过监管规则。\n"
        f"请根据以下合规条款，生成 {k} 个不合规问题，使用投资者口吻表达。\n"
        f"每个问题独立成句，不要包含任何编号、标点或特殊符号。\n\n"
        f"【合规条款】：{clause}"
    )

# 调用模型
async def ask_model(prompt, client):
    try:
        completion = await client.batch_chat.completions.create(
            model=MODEL_ID,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=512,
            temperature=0.9
        )
        response = completion.choices[0].message.content.strip()
        return [q.strip() for q in response.split("\n") if q.strip()]
    except Exception as e:
        print(f"❌ 请求失败: {e}")
        return []

# 主流程
async def process_only_questions(input_path, output_path):
    clauses = extract_clauses(input_path)
    all_questions = []
    
    # 直接创建客户端实例，无需任何配置
    async with AsyncArk() as client:
        # 并发处理每个条款（SDK 自动读取环境变量或构造函数中的 API 密钥）
        for clause in tqdm(clauses, desc="生成问题中", unit="条"):
            prompt = build_prompt(clause, QUESTION_PER_CLAUSE)
            questions = await ask_model(prompt, client)
            all_questions.extend(questions)
    
    # 写入Word文件
    doc = Document()
    doc.add_heading("不合规投资者提问汇总", level=1)
    for q in all_questions:
        doc.add_paragraph(q, style="List Bullet")
    doc.save(output_path)
    print(f"✅ 已保存至: {output_path}")

# 启动程序
if __name__ == "__main__":
    # 将 API 密钥设置为环境变量（SDK 会自动读取）
    os.environ["ARK_API_KEY"] = API_KEY
    
    config = {
        "input_path": "raw_data/客户回访管理办法.docx",
        "output_path": "客户回访不合规问题.docx",
    }

    start = time.perf_counter()
    asyncio.run(process_only_questions(**config))
    print(f"🎉 全部完成，总耗时: {time.perf_counter() - start:.2f} 秒")